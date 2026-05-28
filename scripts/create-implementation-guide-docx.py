#!/usr/bin/env python3
from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
MD_PATH = ROOT / "Server_101_ISO_Implementation_Guide.md"
DOCX_PATH = ROOT / "Server_101_ISO_Implementation_Guide.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(table, top=80, start=120, bottom=80, end=120):
    tbl_pr = table._tbl.tblPr
    tbl_cell_mar = tbl_pr.find(qn("w:tblCellMar"))
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tbl_cell_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tbl_cell_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")
    set_cell_margins(table)
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))


def set_run_font(run, name="Calibri", size=11, color="000000", bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold


def set_style_font(style, name="Calibri", size=11, color="000000", bold=False):
    style.font.name = name
    style._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold


def configure_styles(doc):
    styles = doc.styles

    normal = styles["Normal"]
    set_style_font(normal, "Calibri", 11, "000000")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    title = styles["Title"]
    set_style_font(title, "Calibri", 24, "0B2545", True)
    title.paragraph_format.space_after = Pt(6)

    subtitle = styles["Subtitle"]
    set_style_font(subtitle, "Calibri", 11, "555555")
    subtitle.paragraph_format.space_after = Pt(14)

    for style_name, size, color, before, after in (
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ):
        style = styles[style_name]
        set_style_font(style, "Calibri", size, color, True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25
        style.paragraph_format.keep_with_next = True

    for list_style in ("List Bullet", "List Number"):
        style = styles[list_style]
        set_style_font(style, "Calibri", 11, "000000")
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def add_code_paragraph(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.1)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1.05
    run = p.add_run(text)
    set_run_font(run, "Courier New", 9.5, "0B2545")
    return p


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    set_run_font(r, "Calibri", 11, "1F3A5F", True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    set_run_font(r2, "Calibri", 10.5, "000000")
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)


def add_key_value_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_width(table, [2700, 6660])
    header = table.rows[0].cells
    header[0].text = "Item"
    header[1].text = "Value"
    for cell in header:
        set_cell_shading(cell, "E8EEF5")
        for p in cell.paragraphs:
            for r in p.runs:
                set_run_font(r, bold=True, color="1F4D78")
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
    for row in table.rows[1:]:
        for idx, cell in enumerate(row.cells):
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for r in p.runs:
                    set_run_font(r, color="0B2545" if idx == 0 else "000000", bold=(idx == 0))
    doc.add_paragraph()


def parse_inline_code(paragraph, text):
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, "Courier New", 9.5, "0B2545")
        else:
            run = paragraph.add_run(part)
            set_run_font(run)


def build_doc():
    md = MD_PATH.read_text()
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_styles(doc)

    header = section.header.paragraphs[0]
    header.text = "Server 101 Offline ISO"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        set_run_font(run, "Calibri", 9, "555555")

    footer = section.footer.paragraphs[0]
    footer.text = "Implementation Guide"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        set_run_font(run, "Calibri", 9, "555555")

    title = doc.add_paragraph(style="Title")
    title.add_run("Server 101 Offline ISO Implementation Guide")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run("Field guide for installing the s101 offline Debian server at site 2")

    add_callout(
        doc,
        "Use this guide with the S101_OFFLINE_NEW DVD.",
        "The site 2 server does not need internet access. Packages, installer boot dependencies, and the Server 101 service payload are included on the disc.",
    )

    target_rows = [
        ("Hostname", "s101"),
        ("FQDN", "s101.top.demosdnx.net"),
        ("IP address", "<SERVER_IP>/<PREFIX_LENGTH>"),
        ("Gateway", "<GATEWAY_IP>"),
        ("DNS", "127.0.0.1, <GATEWAY_IP>"),
        ("Primary user", "autoadmin / <CHANGE_ME_PASSWORD>"),
        ("Root password", "<CHANGE_ME_PASSWORD>"),
    ]

    in_code = False
    code_lines = []
    current_heading = ""
    skip_next_target_block = False

    for raw in md.splitlines()[2:]:
        line = raw.rstrip()
        if line.startswith("```"):
            if in_code:
                text = "\n".join(code_lines).strip()
                if text:
                    add_code_paragraph(doc, text)
                code_lines = []
                in_code = False
            else:
                in_code = True
            continue
        if in_code:
            code_lines.append(line)
            continue
        if not line:
            continue
        if line.startswith("## "):
            current_heading = line[3:]
            doc.add_heading(current_heading, level=1)
            if current_heading == "Target Server Requirements":
                add_key_value_table(doc, target_rows)
                skip_next_target_block = True
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:], level=2)
            continue
        if skip_next_target_block and line.startswith("- "):
            continue
        if skip_next_target_block and line.startswith("Default credentials"):
            skip_next_target_block = False
        if re.match(r"^\d+\. ", line):
            p = doc.add_paragraph(style="List Number")
            parse_inline_code(p, re.sub(r"^\d+\. ", "", line))
            continue
        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            parse_inline_code(p, line[2:])
            continue
        p = doc.add_paragraph()
        parse_inline_code(p, line)

    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build_doc()
